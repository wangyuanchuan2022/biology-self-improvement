#!/usr/bin/env python3
"""
自动测试脚本
测试test_question.docx和test_question.pdf文件
"""
import os
import sys
from pathlib import Path

# 添加项目根目录到Python路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from document_processor import DocumentProcessor
from utils.config import get_config


def check_test_files():
    """检查测试文件是否存在"""
    print("检查测试文件...")

    test_files = {
        "test_question.docx": False,
        "test_question.pdf": False
    }

    for file_name in test_files:
        file_path = project_root / file_name
        if file_path.exists():
            test_files[file_name] = True
            print(f"  ✓ 找到 {file_name}")
        else:
            print(f"  ✗ 未找到 {file_name}")

    return test_files


def create_test_files():
    """创建缺失的测试文件"""
    print("\n创建缺失的测试文件...")

    # 创建Word文档
    if not (project_root / "test_question.docx").exists():
        try:
            from docx import Document

            doc = Document()
            doc.add_heading('生物学测试试题', 0)
            doc.add_paragraph('1. 请描述细胞有丝分裂的主要过程。')
            doc.add_paragraph('填空: 有丝分裂分为____、____、____、____四个时期。')
            doc.add_paragraph('2. 比较有丝分裂和减数分裂的主要区别。')
            doc.add_heading('评分标准', 1)
            doc.add_paragraph('1. 准确描述有丝分裂各时期特征（3分）')
            doc.add_paragraph('2. 正确使用专业术语（2分）')
            doc.add_paragraph('3. 比较分析全面准确（5分）')

            doc_path = project_root / "test_question.docx"
            doc.save(doc_path)
            print(f"  ✓ 创建 test_question.docx")
        except ImportError:
            print("  ✗ 无法创建Word文档: python-docx未安装")

    # 创建PDF文档
    if not (project_root / "test_question.pdf").exists():
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.pdfgen import canvas

            pdf_path = project_root / "test_question.pdf"
            c = canvas.Canvas(str(pdf_path), pagesize=letter)
            c.setFont("Helvetica", 12)

            # 添加内容
            c.drawString(100, 750, "生物学测试试题 - PDF版本")
            c.drawString(100, 730, "1. 请描述细胞有丝分裂的主要过程。")
            c.drawString(100, 710, "填空: 有丝分裂分为____、____、____、____四个时期。")
            c.drawString(100, 690, "2. 比较有丝分裂和减数分裂的主要区别。")
            c.drawString(100, 670, " ")
            c.drawString(100, 650, "评分标准")
            c.drawString(100, 630, "1. 准确描述有丝分裂各时期特征（3分）")
            c.drawString(100, 610, "2. 正确使用专业术语（2分）")
            c.drawString(100, 590, "3. 比较分析全面准确（5分）")

            c.save()
            print(f"  ✓ 创建 test_question.pdf")
        except ImportError:
            print("  ✗ 无法创建PDF文档: reportlab未安装")


def test_document_processing():
    """测试文档处理"""
    print("\n测试文档处理...")

    # 获取配置
    try:
        config = get_config()
    except Exception as e:
        print(f"配置加载失败: {e}")
        config = None

    # 创建处理器
    processor_config = {
        "ollama_base_url": config.get("ollama.base_url") if config else "http://localhost:11434",
        "ollama_model": config.get("ollama.model") if config else "qwen3-vl:8b",
        "enable_image_descriptions": False,  # 测试时不启用图片描述
        "pdf_dpi": 150
    }

    processor = DocumentProcessor(config=processor_config)

    # 测试文件列表
    test_files = []
    for file_name in ["test_question.docx", "test_question.pdf"]:
        file_path = project_root / file_name
        if file_path.exists():
            test_files.append(str(file_path))

    if not test_files:
        print("无测试文件可用")
        return False

    print(f"将处理 {len(test_files)} 个文件:")
    for file_path in test_files:
        print(f"  - {os.path.basename(file_path)}")

    # 处理文档
    try:
        results = processor.process_documents(test_files, output_dir="test_output")

        # 打印结果
        print(f"\n处理结果:")
        print(f"  成功: {len(results['successful'])}")
        print(f"  失败: {len(results['failed'])}")

        if results['failed']:
            print("\n失败详情:")
            for fail in results['failed']:
                print(f"  - {os.path.basename(fail['file'])}: {fail['error'][:100]}...")

        # 打印统计信息
        processor.print_statistics()

        # 检查输出文件
        output_dir = Path("test_output")
        if output_dir.exists():
            output_files = list(output_dir.glob("*_processed.txt"))
            if output_files:
                print(f"\n输出文件:")
                for output_file in output_files:
                    print(f"  - {output_file.name}")

        return len(results['failed']) == 0

    except Exception as e:
        print(f"处理失败: {e}")
        import traceback
        traceback.print_exc()
        return False


def main():
    """主函数"""
    print("=" * 60)
    print("自动测试脚本")
    print("=" * 60)

    # 检查测试文件
    test_files = check_test_files()

    # 创建缺失的测试文件
    if not all(test_files.values()):
        create_test_files()
        test_files = check_test_files()

    # 运行文档处理测试
    success = test_document_processing()

    print("\n" + "=" * 60)
    if success:
        print("测试通过! ✓")
        return 0
    else:
        print("测试失败! ✗")
        return 1


if __name__ == "__main__":
    sys.exit(main())