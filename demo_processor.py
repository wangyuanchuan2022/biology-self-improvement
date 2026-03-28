#!/usr/bin/env python3
"""
文档处理器演示脚本
展示如何使用文档处理器模块
"""

import os
import sys
from pathlib import Path

# 添加项目根目录到Python路径
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from document_processor import DocumentProcessor
from utils.config import init_config


def create_test_document():
    """创建测试Word文档"""
    try:
        from docx import Document
    except ImportError:
        print("请先安装 python-docx: pip install python-docx")
        return None

    # 创建测试文档
    doc = Document()

    # 添加标题
    doc.add_heading('生物学测试试题', 0)

    # 添加试题内容
    doc.add_paragraph('1. 请描述细胞有丝分裂的主要过程。')

    doc.add_paragraph('如图所示为细胞有丝分裂的示意图（图1）。')

    doc.add_paragraph('2. 比较有丝分裂和减数分裂的主要区别。')

    # 添加评分标准部分
    doc.add_heading('评分标准', 1)
    doc.add_paragraph('1. 准确描述有丝分裂各时期特征（3分）')
    doc.add_paragraph('2. 正确使用专业术语（2分）')
    doc.add_paragraph('3. 比较分析全面准确（5分）')

    # 保存文档
    test_doc_path = project_root / "test_question.docx"
    doc.save(test_doc_path)

    print(f"测试文档已创建: {test_doc_path}")
    return str(test_doc_path)


def create_test_pdf():
    """创建测试PDF文档"""
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas
    except ImportError:
        print("请先安装 reportlab: pip install reportlab")
        return None

    # 创建测试PDF
    test_pdf_path = project_root / "test_question.pdf"
    c = canvas.Canvas(str(test_pdf_path), pagesize=letter)
    c.setFont("Helvetica", 12)

    # 添加内容
    c.drawString(100, 750, "生物学测试试题 - PDF版本")
    c.drawString(100, 730, "1. 请描述细胞有丝分裂的主要过程。")
    c.drawString(100, 710, " ")
    c.drawString(100, 690, "如图所示为细胞有丝分裂的示意图（图1）。")
    c.drawString(100, 670, "请详细描述图中的各时期特征。")
    c.drawString(100, 650, " ")
    c.drawString(100, 630, "2. 比较有丝分裂和减数分裂的主要区别。")
    c.drawString(100, 610, " ")
    c.drawString(100, 590, "评分标准")
    c.drawString(100, 570, "1. 准确描述有丝分裂各时期特征（3分）")
    c.drawString(100, 550, "2. 正确使用专业术语（2分）")
    c.drawString(100, 530, "3. 比较分析全面准确（5分）")

    c.save()

    print(f"测试PDF已创建: {test_pdf_path}")
    return str(test_pdf_path)


def demonstrate_processor():
    """演示文档处理器功能"""
    print("=" * 60)
    print("文档处理器演示")
    print("=" * 60)

    # 初始化配置
    try:
        config = init_config()
    except ValueError as e:
        print(f"配置初始化失败: {e}")
        print("请检查环境变量或创建 .env 文件")
        return

    # 创建或查找测试文档
    test_docs = []

    # 检查现有测试文档
    possible_docs = [
        "test_question.docx",
        "test_question.pdf",
        "sample_question.docx",
        "sample_question.pdf",
        "生物试题示例.docx",
        "生物试题示例.pdf"
    ]

    for doc_name in possible_docs:
        doc_path = project_root / doc_name
        if doc_path.exists():
            test_docs.append(str(doc_path))
            print(f"找到测试文档: {doc_path}")

    # 确保至少有一个Word文档和一个PDF文档
    has_word = any(doc.endswith('.docx') for doc in test_docs)
    has_pdf = any(doc.endswith('.pdf') for doc in test_docs)

    if not has_word:
        print("未找到Word测试文档，创建测试Word文档...")
        test_doc = create_test_document()
        if test_doc:
            test_docs.append(test_doc)
            has_word = True

    if not has_pdf:
        print("未找到PDF测试文档，创建测试PDF文档...")
        test_pdf = create_test_pdf()
        if test_pdf:
            test_docs.append(test_pdf)
            has_pdf = True

    if not test_docs:
        print("无法创建或找到测试文档，演示终止")
        return

    print(f"\n将处理 {len(test_docs)} 个测试文档:")
    for doc in test_docs:
        print(f"  - {os.path.basename(doc)}")

    # 创建文档处理器
    print("\n创建文档处理器...")
    processor = DocumentProcessor(
        config={
            "ollama_base_url": config.get("ollama.base_url"),
            "ollama_model": config.get("ollama.model"),
            "enable_image_descriptions": config.get("document_processor.enable_image_descriptions"),
            "pdf_dpi": config.get("document_processor.pdf_dpi", 150)
        }
    )

    # 处理文档
    print(f"\n处理 {len(test_docs)} 个文档...")
    try:
        results = processor.process_documents(
            test_docs,
            output_dir=config.get("document_processor.output_dir")
        )

        # 显示统计信息
        processor.print_statistics()

        # 显示成功结果
        if results["successful"]:
            print(f"\n成功处理 {len(results['successful'])} 个文档:")
            for success in results["successful"]:
                print(f"  ✓ {os.path.basename(success['file'])}")

        # 显示失败结果
        if results["failed"]:
            print(f"\n处理失败 {len(results['failed'])} 个文档:")
            for fail in results["failed"]:
                print(f"  ✗ {os.path.basename(fail['file'])}: {fail['error'][:100]}...")

        # 显示输出文件位置
        output_dir = config.get("document_processor.output_dir")
        if output_dir and os.path.exists(output_dir):
            output_files = list(Path(output_dir).glob("*_processed.txt"))
            if output_files:
                print(f"\n输出文件位置: {output_dir}")
                for output_file in output_files:
                    print(f"  📄 {output_file.name}")

    except KeyboardInterrupt:
        print("\n演示被用户中断")
    except Exception as e:
        print(f"\n演示过程中发生错误: {e}")
        import traceback
        traceback.print_exc()

    print("\n" + "=" * 60)
    print("演示结束")
    print("=" * 60)


def check_ollama():
    """检查Ollama服务"""
    import requests

    print("\n检查Ollama服务...")

    config = init_config(".env")
    ollama_url = config.get("ollama.base_url", "http://localhost:11434")
    model = config.get("ollama.model", "qwen3-vl:8b")

    try:
        # 检查服务是否运行
        response = requests.get(f"{ollama_url}/api/tags", timeout=5)
        if response.status_code == 200:
            models = response.json().get("models", [])
            model_names = [m.get("name") for m in models]

            if model in model_names:
                print(f"  ✓ Ollama服务运行正常，模型 {model} 可用")
            else:
                print(f"  ⚠ Ollama服务运行正常，但模型 {model} 未找到")
                print(f"    可用模型: {', '.join(model_names[:3])}{'...' if len(model_names) > 3 else ''}")
                print(f"    请运行: ollama pull {model}")
                return False

            return True
        else:
            print(f"  ✗ Ollama服务异常: {response.status_code}")
            return False

    except requests.exceptions.ConnectionError:
        print("  ✗ 无法连接到Ollama服务")
        print("    请确保Ollama已安装并运行: https://ollama.com/")
        return False
    except Exception as e:
        print(f"  ✗ 检查Ollama服务时出错: {e}")
        return False


def main():
    """主函数"""
    print("生物学习智能体 - 文档处理器演示")
    print()

    # 检查Ollama服务
    if not check_ollama():
        print("\nOllama服务不可用，图片描述功能将无法使用")
        print("可以继续演示文档提取和格式化功能")
        print("按 Enter 继续，或按 Ctrl+C 退出")
        try:
            input()
        except KeyboardInterrupt:
            print("\n演示取消")
            return

    # 执行演示
    demonstrate_processor()

    # 清理建议
    print("\n" + "=" * 60)
    print("下一步建议:")
    print("1. 使用真实试题文档测试")
    print("2. 修改 .env 文件配置Deepseek API密钥")
    print("3. 查看 processed_output/ 目录下的处理结果")
    print("4. 运行测试: pytest tests/")
    print("=" * 60)


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n演示被用户中断")
    except Exception as e:
        print(f"演示过程中发生未预期错误: {e}")
        import traceback
        traceback.print_exc()