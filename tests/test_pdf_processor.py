#!/usr/bin/env python3
"""
测试PDF处理器功能
"""
import os
import sys
import tempfile
from pathlib import Path

# 添加项目根目录到Python路径
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from document_processor import PDFExtractor, DocumentProcessor
from utils.config import get_config


def test_pdf_extractor():
    """测试PDF提取器"""
    print("测试PDF提取器...")

    try:
        # 创建测试PDF文件
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            pdf_path = tmp.name

            # 创建简单的PDF
            c = canvas.Canvas(pdf_path, pagesize=letter)
            c.setFont("Helvetica", 12)
            c.drawString(100, 750, "测试PDF文档")
            c.drawString(100, 730, "1. 请描述细胞有丝分裂的主要过程。")
            c.drawString(100, 710, "填空: 有丝分裂分为____、____、____、____四个时期。")
            c.drawString(100, 690, "2. 比较有丝分裂和减数分裂的主要区别。")
            c.save()

            print(f"创建测试PDF: {pdf_path}")

            # 测试PDF提取器
            extractor = PDFExtractor(dpi=150)
            content = extractor.extract(pdf_path)

            print(f"提取结果: {len(content.images)} 页, {len(content.text_blocks)} 个文本块")

            # 验证结果
            assert len(content.images) == 1, f"应提取1页，实际提取{len(content.images)}页"
            assert content.metadata["source_type"] == "pdf", "元数据中应标记为pdf"

            print("✓ PDF提取器测试通过")

            return True

    except ImportError as e:
        print(f"依赖缺失: {e}")
        return False
    except Exception as e:
        print(f"测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        # 清理临时文件
        if 'pdf_path' in locals() and os.path.exists(pdf_path):
            os.unlink(pdf_path)


def test_document_processor_pdf():
    """测试文档处理器的PDF支持"""
    print("\n测试文档处理器的PDF支持...")

    try:
        # 创建测试PDF文件
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            pdf_path = tmp.name

            # 创建简单的PDF
            c = canvas.Canvas(pdf_path, pagesize=letter)
            c.setFont("Helvetica", 12)
            c.drawString(100, 750, "生物学测试试题")
            c.drawString(100, 730, "1. 请描述细胞有丝分裂的主要过程。")
            c.drawString(100, 710, "填空: 有丝分裂分为____、____、____、____四个时期。")
            c.drawString(100, 690, "2. 比较有丝分裂和减数分裂的主要区别。")
            c.drawString(100, 650, "评分标准")
            c.drawString(100, 630, "1. 准确描述有丝分裂各时期特征（3分）")
            c.drawString(100, 610, "2. 正确使用专业术语（2分）")
            c.drawString(100, 590, "3. 比较分析全面准确（5分）")
            c.save()

            print(f"创建测试PDF: {pdf_path}")

            # 测试文档处理器
            config = get_config()
            processor = DocumentProcessor(
                config={
                    "ollama_base_url": config.get("ollama.base_url"),
                    "ollama_model": config.get("ollama.model"),
                    "enable_image_descriptions": False,  # 测试时不启用图片描述
                    "pdf_dpi": 150
                }
            )

            # 处理PDF
            agent_input = processor.process_document(pdf_path)

            print(f"处理完成: 试题文本长度={len(agent_input.question_text)}")
            print(f"评分标准: {agent_input.scoring_standard[:100]}...")

            # 验证结果
            assert agent_input.question_text, "试题文本不应为空"
            assert agent_input.scoring_standard, "评分标准不应为空"
            assert "有丝分裂" in agent_input.question_text, "试题文本应包含'有丝分裂'"

            print("✓ 文档处理器PDF支持测试通过")

            return True

    except ImportError as e:
        print(f"依赖缺失: {e}")
        return False
    except Exception as e:
        print(f"测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        # 清理临时文件
        if 'pdf_path' in locals() and os.path.exists(pdf_path):
            os.unlink(pdf_path)


def test_word_and_pdf_mixed():
    """测试混合格式文档处理"""
    print("\n测试混合格式文档处理...")

    try:
        # 创建临时Word和PDF文件
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        temp_files = []

        # 创建Word文档
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            docx_path = tmp.name
            doc = Document()
            doc.add_heading('Word测试试题', 0)
            doc.add_paragraph('1. 请描述光合作用的主要过程。')
            doc.save(docx_path)
            temp_files.append(docx_path)
            print(f"创建测试Word文档: {docx_path}")

        # 创建PDF文档
        with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
            pdf_path = tmp.name
            c = canvas.Canvas(pdf_path, pagesize=letter)
            c.setFont("Helvetica", 12)
            c.drawString(100, 750, "PDF测试试题")
            c.drawString(100, 730, "1. 请描述呼吸作用的主要过程。")
            c.save()
            temp_files.append(pdf_path)
            print(f"创建测试PDF文档: {pdf_path}")

        # 测试混合处理
        config = get_config()
        processor = DocumentProcessor(
            config={
                "ollama_base_url": config.get("ollama.base_url"),
                "ollama_model": config.get("ollama.model"),
                "enable_image_descriptions": False,
                "pdf_dpi": 150
            }
        )

        results = processor.process_documents(temp_files, output_dir=None)

        print(f"混合处理结果: 成功={len(results['successful'])}, 失败={len(results['failed'])}")

        # 验证结果
        assert len(results['successful']) == 2, f"应成功处理2个文档，实际成功{len(results['successful'])}个"
        assert len(results['failed']) == 0, f"应无失败，实际失败{len(results['failed'])}个"

        print("✓ 混合格式文档处理测试通过")

        return True

    except ImportError as e:
        print(f"依赖缺失: {e}")
        return False
    except Exception as e:
        print(f"测试失败: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        # 清理临时文件
        for file_path in temp_files:
            if os.path.exists(file_path):
                os.unlink(file_path)


def main():
    """主测试函数"""
    print("=" * 60)
    print("PDF处理器功能测试")
    print("=" * 60)

    # 运行测试
    tests = [
        test_pdf_extractor,
        test_document_processor_pdf,
        test_word_and_pdf_mixed
    ]

    passed = 0
    failed = 0

    for test_func in tests:
        try:
            if test_func():
                passed += 1
            else:
                failed += 1
        except Exception as e:
            print(f"测试执行异常: {e}")
            failed += 1

    print("\n" + "=" * 60)
    print(f"测试结果: 通过 {passed}/{len(tests)}, 失败 {failed}/{len(tests)}")

    if failed == 0:
        print("所有测试通过! ✓")
        return 0
    else:
        print(f"{failed} 个测试失败! ✗")
        return 1


if __name__ == "__main__":
    sys.exit(main())