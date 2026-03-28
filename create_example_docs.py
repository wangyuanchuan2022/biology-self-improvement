#!/usr/bin/env python3
"""
创建示例Word文档用于测试
"""

import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def create_example_documents():
    """创建示例文档"""
    project_root = Path(__file__).parent

    # 文档1：简单试题
    doc1 = Document()
    doc1.add_heading('生物学测试试题 - 示例 1', 0)

    # 试题内容
    p1 = doc1.add_paragraph()
    p1.add_run('1. 请描述细胞有丝分裂的主要过程。').bold = True

    doc1.add_paragraph('如图所示为细胞有丝分裂的示意图（图1）。')
    doc1.add_paragraph('请详细描述图中的各时期特征。')

    # 评分标准
    doc1.add_heading('评分标准', 1)
    doc1.add_paragraph('1. 准确描述有丝分裂各时期特征（3分）')
   1.add_paragraph('2. 正确使用专业术语（2分）')
    doc1.add_paragraph('3. 比较分析全面准确（5分）')

    # 保存文档

    doc1_path = project_root / "test_question.docx"
    doc1.save(doc1_path)

    print(f"示例文档1已创建: {doc1_path}")

    # 文档2：包含表格的试题

    doc2= Document()
    doc2.add_heading('生物学测试试题 - 示例 2（含表格）', 0)

    # 试题内容
    doc2.add_paragraph('2. 根据下表数据，分析不同温度对酶活性的影响。')

    # 添加表格
   table = doc2.add_table(rows=5, cols=3)
    table.style='TableGrid'

    # 设置表头
    header_cells=table.rows[0].cells
   header_cells[0].text='温度(℃)'
    header_cells[1].text='反应速度(mmol/min)'
    header_cells[2].text='相对活性(%)'

   # 添加数据
   data = [
        (20, 0.5, 25),
        (30, 1.2, 60),
        (40, 2.0, 100),
        (50, 0.8, 40),
    ]

   for i,(temp,rate,activity) in enumerate(data,1):
        row_cells = table.rows[i].cells
        row_cells[0].text = str(temp)
        row_cells[1].text = str(rate)
        row_cells[2].text = str(activity)

    # 保存文档

   doc2_path = project_root / "sample_question.docx"
    doc2.save(doc2_path)

    print(f"示例文档2已创建: {doc2_path}")

    return [str(doc1_path), str(doc2_path)]



if __name__ == "__main__":
    print("创建示例Word文档...")
    docs=create_example_documents()

   print(f"\n已创建{len(docs)}个示例文档:")
    for doc in docs:
        print(f"  - {doc}")

   print("\n使用建议:")
   print("1. 运行 demo_processor.py 测试文档处理功能")
   print("2. 使用真实试题文档进行更全面的测试")
   print("3. 修改 document_processor/ 中的代码以适应实际需求")