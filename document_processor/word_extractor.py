"""
Word文档提取器
使用python-docx提取Word文档中的图片和文字
"""

import io
import os
from typing import List, Optional, Tuple
from docx import Document
from docx.document import Document as DocxDocument
from docx.text.paragraph import Paragraph
from docx.shared import Inches
import zipfile

from .models import DocumentImage, TextBlock, DocumentContent, ImageType


class WordExtractor:
    """Word文档提取器"""

    def __init__(self):
        self.document: Optional[DocxDocument] = None

    def extract(self, docx_path: str) -> DocumentContent:
        """
        提取Word文档内容

        Args:
            docx_path: Word文档路径

        Returns:
            DocumentContent: 文档内容
        """
        if not os.path.exists(docx_path):
            raise FileNotFoundError(f"文档不存在: {docx_path}")

        # 加载Word文档
        self.document = Document(docx_path)

        # 提取图片和文本
        images = self._extract_images(docx_path)
        text_blocks = self._extract_text_blocks()

        # 创建文档内容对象
        content = DocumentContent(
            images=images,
            text_blocks=text_blocks,
            metadata={
                "file_path": docx_path,
                "file_name": os.path.basename(docx_path),
                "total_images": len(images),
                "total_text_blocks": len(text_blocks)
            }
        )

        return content

    def _extract_images(self, docx_path: str) -> List[DocumentImage]:
        """
        从Word文档中提取图片

        Args:
            docx_path: Word文档路径

        Returns:
            图片列表
        """
        images = []

        # Word文档本质上是ZIP文件
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # 查找媒体文件
            media_files = [f for f in docx_zip.namelist()
                          if f.startswith('word/media/')]

            for i, media_path in enumerate(media_files):
                # 提取图片数据
                image_data = docx_zip.read(media_path)

                # 确定图片格式
                image_format = self._get_image_format(media_path)

                # 创建图片对象
                image = DocumentImage(
                    image_id=f"img_{i+1:03d}",
                    image_data=image_data,
                    image_format=image_format,
                    image_type=self._detect_image_type(media_path, image_data)
                )

                images.append(image)

        return images

    def _extract_text_blocks(self) -> List[TextBlock]:
        """
        提取文本块

        Returns:
            文本块列表
        """
        text_blocks = []

        if not self.document:
            return text_blocks

        # 提取段落文本
        for paragraph in self.document.paragraphs:
            text = paragraph.text.strip()
            if text:  # 跳过空段落
                # 简单样式检测
                style = self._detect_paragraph_style(paragraph)

                text_block = TextBlock(
                    text=text,
                    style=style
                )
                text_blocks.append(text_block)

        # 提取表格文本
        for table in self.document.tables:
            table_text = self._extract_table_text(table)
            if table_text:
                text_block = TextBlock(
                    text=table_text,
                    style="table"
                )
                text_blocks.append(text_block)

        return text_blocks

    def _get_image_format(self, media_path: str) -> str:
        """
        根据文件路径获取图片格式

        Args:
            media_path: 媒体文件路径

        Returns:
            图片格式 (jpg, png, 等)
        """
        # 从文件扩展名提取格式
        ext = os.path.splitext(media_path)[1].lower()
        if ext in ['.jpg', '.jpeg']:
            return 'jpeg'
        elif ext == '.png':
            return 'png'
        elif ext == '.gif':
            return 'gif'
        elif ext in ['.bmp', '.dib']:
            return 'bmp'
        elif ext == '.tiff':
            return 'tiff'
        else:
            return 'unknown'

    def _detect_image_type(self, media_path: str, image_data: bytes) -> ImageType:
        """
        检测图片类型（基于文件名和内容）

        Args:
            media_path: 媒体文件路径
            image_data: 图片数据

        Returns:
            图片类型
        """
        # 基于文件名的简单检测
        filename = os.path.basename(media_path).lower()

        # 关键词检测（可根据实际情况扩展）
        if any(keyword in filename for keyword in ['实验', 'experiment', 'exp']):
            return ImageType.EXPERIMENT
        elif any(keyword in filename for keyword in ['图表', 'chart', 'graph', 'figure']):
            return ImageType.CHART
        elif any(keyword in filename for keyword in ['流程图', 'diagram', 'flow']):
            return ImageType.DIAGRAM
        elif any(keyword in filename for keyword in ['显微', 'microscope', '细胞']):
            return ImageType.MICROSCOPE
        elif any(keyword in filename for keyword in ['结构', 'structure']):
            return ImageType.STRUCTURE
        elif any(keyword in filename for keyword in ['表格', 'table']):
            return ImageType.TABLE
        else:
            return ImageType.OTHER

    def _detect_paragraph_style(self, paragraph: Paragraph) -> Optional[str]:
        """
        检测段落样式

        Args:
            paragraph: 段落对象

        Returns:
            样式名称
        """
        style_name = paragraph.style.name.lower()

        # 简单样式分类
        if 'heading' in style_name or '标题' in style_name:
            return 'heading'
        elif 'title' in style_name:
            return 'title'
        elif 'caption' in style_name or '题注' in style_name:
            return 'caption'
        else:
            return 'normal'

    def _extract_table_text(self, table) -> str:
        """
        提取表格文本

        Args:
            table: 表格对象

        Returns:
            表格文本表示
        """
        table_text = []

        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)

            if row_text:
                table_text.append(" | ".join(row_text))

        return "\n".join(table_text)


# 使用示例
if __name__ == "__main__":
    # 测试代码
    extractor = WordExtractor()
    try:
        # 使用一个示例文档路径
        test_doc = "示例文档.docx"  # 需要实际存在的文档
        if os.path.exists(test_doc):
            content = extractor.extract(test_doc)
            print(f"提取成功: {len(content.images)} 张图片, {len(content.text_blocks)} 个文本块")
        else:
            print("测试文档不存在，跳过测试")
    except Exception as e:
        print(f"提取失败: {e}")